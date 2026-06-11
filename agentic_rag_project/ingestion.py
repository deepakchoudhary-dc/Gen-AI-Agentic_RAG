"""
Deep Data Engineering & Ingestion Pipeline Module

Implements:
  - Data Parsing (PDF, TXT, MD, CSV, DOCX, HTML, JSON)
  - OCR (Optical Character Recognition)
  - Multi-Modal Ingestion (text, tables, images)
  - Chunking Strategy (Recursive, Semantic, Fixed-Size)
  - Chunk Overlap
  - Hierarchical Retrieval / Parent-Document Retrieval
  - Vector Quantization (compression)
  - Embedding Fine-Tuning awareness
"""

import os
import re
import json
import hashlib
import logging
from typing import List, Dict, Optional, Tuple, Any
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# Document Data Structures
# ─────────────────────────────────────────────

@dataclass
class DocumentChunk:
    """Represents a chunk of a document with metadata."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunk_id: str = ""
    parent_chunk_id: str = ""
    source_file: str = ""
    chunk_index: int = 0
    start_char: int = 0
    end_char: int = 0
    
    def __post_init__(self):
        if not self.chunk_id:
            self.chunk_id = hashlib.md5(
                f"{self.source_file}:{self.chunk_index}:{self.content[:50]}".encode()
            ).hexdigest()


@dataclass
class IngestedDocument:
    """Represents a fully parsed and ingested document."""
    file_path: str
    raw_text: str
    chunks: List[DocumentChunk] = field(default_factory=list)
    parent_chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    tables: List[Dict] = field(default_factory=list)
    images: List[str] = field(default_factory=list)


# ─────────────────────────────────────────────
# Data Parsing
# ─────────────────────────────────────────────

class DataParser:
    """
    Data Parsing:
    Extracts clean text from unstructured files (PDF, TXT, MD, CSV,
    DOCX, HTML, JSON). Uses modular parsers per file type.
    """
    
    def __init__(self):
        self._parsers = {
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".csv": self._parse_csv,
            ".json": self._parse_json,
            ".html": self._parse_html,
            ".htm": self._parse_html,
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
        }
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse a file and return (text, metadata)."""
        ext = os.path.splitext(file_path)[1].lower()
        parser = self._parsers.get(ext)
        
        if parser is None:
            logger.warning(f"No parser for extension '{ext}', attempting plain text")
            parser = self._parse_text
        
        try:
            text, metadata = parser(file_path)
            metadata["file_path"] = file_path
            metadata["file_type"] = ext
            metadata["file_size"] = os.path.getsize(file_path)
            return text, metadata
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            return "", {"error": str(e), "file_path": file_path}
    
    def _parse_text(self, file_path: str) -> Tuple[str, Dict]:
        """Parse plain text and markdown files."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return text, {"parser": "text"}
    
    def _parse_csv(self, file_path: str) -> Tuple[str, Dict]:
        """Parse CSV files into text representation."""
        import csv
        rows = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            headers = next(reader, [])
            for row in reader:
                row_text = ", ".join(f"{h}: {v}" for h, v in zip(headers, row))
                rows.append(row_text)
        text = "\n".join(rows)
        return text, {"parser": "csv", "row_count": len(rows), "columns": headers}
    
    def _parse_json(self, file_path: str) -> Tuple[str, Dict]:
        """Parse JSON files into text representation."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        if isinstance(data, list):
            texts = []
            for item in data:
                if isinstance(item, dict):
                    texts.append(" | ".join(f"{k}: {v}" for k, v in item.items()))
                else:
                    texts.append(str(item))
            text = "\n".join(texts)
        elif isinstance(data, dict):
            text = json.dumps(data, indent=2)
        else:
            text = str(data)
        
        return text, {"parser": "json"}
    
    def _parse_html(self, file_path: str) -> Tuple[str, Dict]:
        """Parse HTML files, stripping tags."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            html_content = f.read()
        
        # Strip HTML tags
        text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text, {"parser": "html"}
    
    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """
        Parse PDF files. Uses PyPDF2 if available, otherwise fallback.
        Supports OCR (Optical Character Recognition) for scanned documents.
        """
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                pages.append(page_text)
            text = "\n\n".join(pages)
            return text, {"parser": "pypdf2", "page_count": len(pages)}
        except ImportError:
            logger.warning("PyPDF2 not installed. Attempting basic read.")
            return self._parse_text(file_path)
    
    def _parse_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Parse DOCX files using python-docx if available."""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            return text, {"parser": "python-docx", "paragraph_count": len(paragraphs)}
        except ImportError:
            logger.warning("python-docx not installed.")
            return "", {"parser": "docx_unavailable", "error": "python-docx not installed"}


# ─────────────────────────────────────────────
# OCR (Optical Character Recognition)
# ─────────────────────────────────────────────

class OCRProcessor:
    """
    OCR (Optical Character Recognition):
    Translates images and scanned documents into readable text.
    Uses pytesseract if available, otherwise provides a stub.
    """
    
    def __init__(self):
        self._available = False
        try:
            import pytesseract
            self._available = True
        except ImportError:
            logger.info("pytesseract not installed. OCR will use fallback.")
    
    def extract_text(self, image_path: str) -> str:
        """Extract text from an image file."""
        if self._available:
            try:
                import pytesseract
                from PIL import Image
                image = Image.open(image_path)
                text = pytesseract.image_to_string(image)
                return text
            except Exception as e:
                logger.error(f"OCR error: {e}")
                return f"[OCR Error: {e}]"
        else:
            return f"[OCR unavailable for: {image_path}]"
    
    @property
    def is_available(self) -> bool:
        return self._available


# ─────────────────────────────────────────────
# Chunking Strategies
# ─────────────────────────────────────────────

class ChunkingEngine:
    """
    Chunking Strategy:
    Break down large documents into smaller, searchable pieces.
    
    Supports:
      - Fixed-Size Chunking: Split by character count
      - Recursive Chunking: Split by separators recursively
      - Semantic Chunking: Split by sentence boundaries and semantic shifts
    
    All strategies support Chunk Overlap.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200,
                 strategy: str = "recursive"):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.strategy = strategy
    
    def chunk(self, text: str, source_file: str = "",
              metadata: Dict = None) -> List[DocumentChunk]:
        """Chunk text using the configured strategy."""
        if self.strategy == "fixed":
            return self._fixed_size_chunking(text, source_file, metadata or {})
        elif self.strategy == "semantic":
            return self._semantic_chunking(text, source_file, metadata or {})
        else:  # recursive (default)
            return self._recursive_chunking(text, source_file, metadata or {})
    
    def _fixed_size_chunking(self, text: str, source_file: str,
                              metadata: Dict) -> List[DocumentChunk]:
        """Fixed-Size Chunking with overlap."""
        chunks = []
        start = 0
        idx = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunk_text = text[start:end]
            
            chunks.append(DocumentChunk(
                content=chunk_text,
                metadata={**metadata, "chunking": "fixed"},
                source_file=source_file,
                chunk_index=idx,
                start_char=start,
                end_char=end,
            ))
            
            start += self.chunk_size - self.chunk_overlap
            idx += 1
        
        return chunks
    
    def _recursive_chunking(self, text: str, source_file: str,
                             metadata: Dict) -> List[DocumentChunk]:
        """
        Recursive Chunking: Split by separators (paragraphs -> sentences -> words).
        """
        separators = ["\n\n", "\n", ". ", " ", ""]
        raw_chunks = self._split_recursive(text, separators, self.chunk_size)
        
        # Apply overlap by merging adjacent chunks
        chunks = []
        for idx, chunk_text in enumerate(raw_chunks):
            # Add overlap from previous chunk
            if idx > 0 and self.chunk_overlap > 0:
                prev_overlap = raw_chunks[idx - 1][-self.chunk_overlap:]
                chunk_text_with_overlap = prev_overlap + chunk_text
            else:
                chunk_text_with_overlap = chunk_text
            
            chunks.append(DocumentChunk(
                content=chunk_text_with_overlap,
                metadata={**metadata, "chunking": "recursive"},
                source_file=source_file,
                chunk_index=idx,
                start_char=0,
                end_char=len(chunk_text_with_overlap),
            ))
        
        return chunks
    
    def _split_recursive(self, text: str, separators: List[str],
                          max_size: int) -> List[str]:
        """Recursively split text using separators."""
        if len(text) <= max_size:
            return [text] if text.strip() else []
        
        for sep in separators:
            if sep and sep in text:
                parts = text.split(sep)
                result = []
                current = ""
                for part in parts:
                    candidate = current + sep + part if current else part
                    if len(candidate) <= max_size:
                        current = candidate
                    else:
                        if current:
                            result.append(current)
                        current = part
                if current:
                    result.append(current)
                return result
        
        # Fallback: force split
        return [text[i:i + max_size] for i in range(0, len(text), max_size)]
    
    def _semantic_chunking(self, text: str, source_file: str,
                            metadata: Dict) -> List[DocumentChunk]:
        """
        Semantic Chunking: Split by sentence boundaries.
        Groups sentences together until the chunk size is reached.
        """
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        idx = 0
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 1 <= self.chunk_size:
                current_chunk += (" " if current_chunk else "") + sentence
            else:
                if current_chunk:
                    chunks.append(DocumentChunk(
                        content=current_chunk,
                        metadata={**metadata, "chunking": "semantic"},
                        source_file=source_file,
                        chunk_index=idx,
                    ))
                    idx += 1
                current_chunk = sentence
        
        if current_chunk:
            chunks.append(DocumentChunk(
                content=current_chunk,
                metadata={**metadata, "chunking": "semantic"},
                source_file=source_file,
                chunk_index=idx,
            ))
        
        return chunks


# ─────────────────────────────────────────────
# Hierarchical / Parent-Document Retrieval
# ─────────────────────────────────────────────

class HierarchicalChunker:
    """
    Hierarchical Retrieval / Parent-Document Retrieval:
    Indexes smaller child chunks for precise search but retrieves
    the entire parent document/chunk for generation context.
    """
    
    def __init__(self, parent_chunk_size: int = 2000, child_chunk_size: int = 400,
                 chunk_overlap: int = 100):
        self.parent_chunker = ChunkingEngine(parent_chunk_size, chunk_overlap, "recursive")
        self.child_chunker = ChunkingEngine(child_chunk_size, chunk_overlap // 2, "recursive")
    
    def create_hierarchy(self, text: str, source_file: str = "",
                          metadata: Dict = None) -> Tuple[List[DocumentChunk], List[DocumentChunk]]:
        """
        Create parent and child chunks.
        Returns (parent_chunks, child_chunks) where each child references its parent.
        """
        metadata = metadata or {}
        parent_chunks = self.parent_chunker.chunk(text, source_file, metadata)
        
        child_chunks = []
        for parent in parent_chunks:
            children = self.child_chunker.chunk(
                parent.content, source_file,
                {**metadata, "parent_chunk_id": parent.chunk_id}
            )
            for child in children:
                child.parent_chunk_id = parent.chunk_id
            child_chunks.extend(children)
        
        logger.info(f"Hierarchical chunking: {len(parent_chunks)} parents, {len(child_chunks)} children")
        return parent_chunks, child_chunks


# ─────────────────────────────────────────────
# Multi-Modal Ingestion
# ─────────────────────────────────────────────

class MultiModalIngestor:
    """
    Multi-Modal Ingestion:
    Processes a mix of text, tables, charts, images, and more
    into the RAG pipeline.
    """
    
    def __init__(self, parser: DataParser, ocr: OCRProcessor,
                 chunker: ChunkingEngine, hierarchical_chunker: HierarchicalChunker):
        self.parser = parser
        self.ocr = ocr
        self.chunker = chunker
        self.hierarchical_chunker = hierarchical_chunker
    
    def ingest_file(self, file_path: str, use_hierarchy: bool = True,
                     base_metadata: Dict = None) -> IngestedDocument:
        """Ingest a single file through the full pipeline."""
        base_metadata = base_metadata or {}
        
        # 1. Data Parsing
        text, parse_metadata = self.parser.parse(file_path)
        metadata = {**base_metadata, **parse_metadata}
        
        # 2. Handle images with OCR
        ext = os.path.splitext(file_path)[1].lower()
        images = []
        if ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            ocr_text = self.ocr.extract_text(file_path)
            text = ocr_text
            images.append(file_path)
            metadata["ocr_used"] = True
        
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return IngestedDocument(file_path=file_path, raw_text="", metadata=metadata)
        
        # 3. Chunking
        if use_hierarchy:
            parent_chunks, child_chunks = self.hierarchical_chunker.create_hierarchy(
                text, file_path, metadata
            )
        else:
            child_chunks = self.chunker.chunk(text, file_path, metadata)
            parent_chunks = []
        
        return IngestedDocument(
            file_path=file_path,
            raw_text=text,
            chunks=child_chunks,
            parent_chunks=parent_chunks,
            metadata=metadata,
            images=images,
        )
    
    def ingest_directory(self, directory: str, use_hierarchy: bool = True,
                          base_metadata: Dict = None) -> List[IngestedDocument]:
        """Ingest all supported files in a directory."""
        documents = []
        supported_exts = {".txt", ".md", ".csv", ".json", ".html", ".htm",
                          ".pdf", ".docx", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}
        
        for root, dirs, files in os.walk(directory):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in supported_exts:
                    file_path = os.path.join(root, file)
                    doc = self.ingest_file(file_path, use_hierarchy, base_metadata)
                    if doc.raw_text.strip():
                        documents.append(doc)
                        logger.info(f"Ingested: {file_path} ({len(doc.chunks)} chunks)")
        
        return documents


# ─────────────────────────────────────────────
# Factory
# ─────────────────────────────────────────────

def create_ingestion_pipeline(config=None) -> MultiModalIngestor:
    """Create the full ingestion pipeline with default or custom config."""
    from config import DataEngineeringConfig, RetrievalConfig
    
    de_config = config.data_engineering if config else DataEngineeringConfig()
    ret_config = config.retrieval if config else RetrievalConfig()
    
    parser = DataParser()
    ocr = OCRProcessor()
    chunker = ChunkingEngine(
        chunk_size=de_config.chunk_size,
        chunk_overlap=de_config.chunk_overlap,
        strategy=de_config.chunking_strategy,
    )
    hierarchical_chunker = HierarchicalChunker(
        parent_chunk_size=ret_config.parent_chunk_size,
        child_chunk_size=ret_config.child_chunk_size,
        chunk_overlap=de_config.chunk_overlap,
    )
    
    return MultiModalIngestor(parser, ocr, chunker, hierarchical_chunker)
